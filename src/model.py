"""
Generate checkout cards for ETHOS Lab library.

Uses a Word file as a template.
"""

import logging
import json
from datetime import datetime
from time import sleep
from configparser import ConfigParser
import requests
import pandas as pd
from docx import Document

logger = logging.getLogger()

DEBUG = True

if DEBUG:
    logger.setLevel(logging.DEBUG)


class CheckoutCard():
    """A checkout card."""

    FILE = 'templates/ethos library checkout card.docx'
    OUTPUTDIR = 'output'

    def __init__(self, author, title, year):
        """Construct the object.

        Parameters:
        author (str): Lastname, Firstname
        title (str):  Book title
        year (int):   Publication year
        """
        self._doc = Document(self.FILE)
        self._author_form = self._doc.paragraphs[0]
        self._title_form = self._doc.paragraphs[1]
        self._year_form = self._doc.paragraphs[2]

        self.author = author
        self.title = title
        self.year = year

        self._generate_id()
        self._populate()

        logging.debug(repr(self))

    def _generate_id(self):
        id_a = self.author.lower().replace(' ', '_')
        id_t = self.title.lower().replace(' ', '_')
        id_y = self.year
        self._id = f"{id_a}-{id_t}-{id_y}"

    def _populate(self):
        """Populate the form."""
        self._author_form.text = self.author
        self._title_form.text = self.title
        self._year_form.text = str(self.year)

    def __str__(self):
        """Return as string.."""
        return f"{self.author}: {self.title} ({self.year})"

    def __repr__(self):
        """Return as representation."""
        cls = self.__class__
        author = self.author
        title = self.title
        year = self.year
        # return f"{self.__class__}: {self.author}; {self.title}; {self.year}"
        return f'{cls}(author="{author}",title="{title}",year={year})'

    def save(self, filename=None):
        """Save the document to a file.

        Parameters:
        filename (str): file name
        """
        if not filename:
            filename = f"{self.OUTPUTDIR}/{self._id}.docx"

        logging.debug("Writing %s", filename)
        self._doc.save(filename)


class CheckoutCardTwoCols(CheckoutCard):
    """A two-colung checkout card."""

    FILE = 'templates/ethos library checkout card-two columns.docx'

    # def __init__(self, author_l, title_l, year_l,
    #             author_r, title_r, year_r):
    def __init__(self, left_data, right_data):
        self._doc = Document(self.FILE)

        self._author_l_form = self._doc.paragraphs[0]
        self._title_l_form = self._doc.paragraphs[1]
        self._year_l_form = self._doc.paragraphs[2]

        # self.author_l = author_l
        # self.title_l = title_l
        # self.year_l = year_l

        self.author_l = ", ".join(left_data['authors'])
        # self.title_l = ". ".join([left_data['title'], left_data.get('subtitle', '')])
        self.title_l = left_data['title']
        try:
            self.year_l = datetime.fromisoformat(left_data['publishedDate']).year
        except ValueError:
            self.year_l = left_data['publishedDate'].split('-')[0]
        except KeyError:
            self.year_l = ' '

        self._author_r_form = self._doc.paragraphs[10].runs[7]
        self._title_r_form = self._doc.paragraphs[11]
        self._year_r_form = self._doc.paragraphs[12]

        # self.author_r = author_r
        # self.title_r = title_r
        # self.year_r = year_r

        self.author_r = ", ".join(right_data['authors'])
        # self.title_r = ". ".join([right_data['title'], right_data.get('subtitle', '')])
        self.title_r = right_data['title']
        try:
            self.year_r = datetime.fromisoformat(right_data['publishedDate']).year
        except ValueError:
            self.year_r = right_data['publishedDate'].split('-')[0]
        except KeyError:
            self.yearr = ' '

        # self._generate_id()
        # self._populate()

    def _populate(self):
        """Populate the form."""
        self._author_l_form.text = self.author_l
        self._title_l_form.text = self.title_l
        self._year_l_form.text = str(self.year_l)

        self._author_r_form.text = self.author_r
        self._title_r_form.text = self.title_r
        self._year_r_form.text = str(self.year_r)

    def __repr__(self):
        left = self.title_l
        right = self.title_r
        return f'{left} – {right}'


class GoogleBooks():
    """A Google Books API representation."""

    ENDPOINT = 'https://www.googleapis.com/books/v1/volumes'
    FIELDS = 'items/volumeInfo(authors,title,subtitle,publishedDate)'

    def __init__(self, configfile):
        """Construct the Google Books API representation."""
        self._config = ConfigParser()
        self._config.read(configfile)
        logging.debug("Made %s", self.__dir__)

    def get_by_isbn(self, isbn, throttle=2):
        """Get bibliographical data based on an ISBN."""
        logging.debug("Querying %s", isbn)
        creds = {'key': self._config['google']['key']}
        query = {'q': 'isbn:' + isbn, 'fields': self.FIELDS}
        params = {**creds, **query}

        logging.debug("Throttling for %s seconds", throttle)
        sleep(throttle)

        logging.debug("Would query %s with %s", self.ENDPOINT, params)
        resp = requests.get(self.ENDPOINT, params=params)
        logging.debug(resp)

        return resp.json()


class Catalogue():
    """A catalogue for ETHOS Lab."""

    JSONFILE = 'ethos-lib.catalogue.json'

    def __init__(self, datafile, resolver):
        """Construct the catalogue representation."""
        logging.debug("Reading catalogue from %s", datafile)
        self._resolver = resolver
        self._barcodes = pd.read_csv(datafile,
                                     dtype={'Barcode': str},
                                     sep=';')
        self.items = {}
        self.load()
        # self.populate()

    def load(self):
        """Load from storage."""
        with open(self.JSONFILE) as prior_storage:
            self.items = json.load(prior_storage)
        logging.debug("Loaded %s items", len(self.items))

    def populate(self, persist=True):
        """Populate the catalogue metadata.

        # Params:
        persist (bool): Persist to storage.
        """

        for barcode in self._barcodes['Barcode']:
            if barcode not in self.items:
                work = self._resolver.get_by_isbn(barcode)
                try:
                    volume = work['items'][0]['volumeInfo']
                    self.items[barcode] = volume
                except KeyError:
                    ...

        if persist:
            self.persist()

    def persist(self):
        """Persist to a file."""
        logging.debug("Persisting %s items", len(self.items))
        with open(self.JSONFILE, 'w') as new_storage:
            serialized_items = json.dumps(self.items, indent=2)
            new_storage.write(serialized_items)


class PublicationWork():
    """A representation of a publication."""

    def __init__(self, author, title, subtitle, year):
        """Construct the object."""
        self.author = author
        self.title = title
        self.subtitle = subtitle
        self.year = year

    def __str__(self):
        """Return as string."""
        return f"{self.author}: {self.title} ({self.year})"
